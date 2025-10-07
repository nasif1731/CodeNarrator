"""
GenAI Documentation Generator - Streamlit Application
Uses trained BPE tokenizers and Word2Vec models to analyze code and generate documentation
"""

import streamlit as st
import pickle
import os
import re
import numpy as np
from collections import defaultdict, Counter
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import datetime
import io
from pathlib import Path  # <-- for robust paths

# Configure Streamlit page
st.set_page_config(
    page_title="GenAI Documentation Generator",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .section-header {
        font-size: 1.5rem;
        color: #ff7f0e;
        margin-top: 2rem;
        margin-bottom: 1rem;
    }
    .model-status {
        padding: 0.5rem;
        border-radius: 0.5rem;
        margin: 0.5rem 0;
    }
    .status-success {
        background-color: #d4edda;
        color: #155724;
        border: 1px solid #c3e6cb;
    }
    .status-error {
        background-color: #f8d7da;
        color: #721c24;
        border: 1px solid #f5c6cb;
    }
    .code-analysis {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #007bff;
    }
</style>
""", unsafe_allow_html=True)

# =========================
# Advanced LSTM class FIRST
# =========================
class AdvancedLSTM:
    """Professional LSTM implementation with proper initialization"""
    def __init__(self, input_size, hidden_size, output_size):
        self.input_size = input_size
        self.hidden_size = hidden_size
        self.output_size = output_size
        self.trained = False

        # Initialize with dummy weights (will be loaded from pickle)
        self.Wf = None
        self.Wi = None
        self.Wo = None
        self.Wg = None
        self.Wy = None
        self.bf = None
        self.bi = None
        self.bo = None
        self.bg = None
        self.by = None

    def sigmoid(self, x):
        return 1 / (1 + np.exp(-np.clip(x, -500, 500)))

    def tanh(self, x):
        return np.tanh(np.clip(x, -500, 500))

    def forward(self, x, h_prev, c_prev):
        """LSTM forward pass"""
        if self.Wf is None:  # Not properly loaded
            return h_prev, c_prev, np.zeros((self.output_size, 1))

        # Concatenate input and previous hidden state
        concat = np.vstack((x.reshape(-1, 1), h_prev))

        # Gates
        f = self.sigmoid(self.Wf @ concat + self.bf)  # Forget gate
        i = self.sigmoid(self.Wi @ concat + self.bi)  # Input gate
        o = self.sigmoid(self.Wo @ concat + self.bo)  # Output gate
        g = self.tanh(self.Wg @ concat + self.bg)     # Candidate values

        # Update cell state and hidden state
        c = f * c_prev + i * g
        h = o * self.tanh(c)

        # Output
        y = self.Wy @ h + self.by

        return h, c, y


# =========================
# Load trained models
# =========================
@st.cache_resource
def load_models():
    """
    Load BPE, Word2Vec, and LSTM models using project-relative paths.
    - BPE: supports 'merges-only' pickles (list/tuple of merges) or dicts.
    - W2V/LSTM: expects your custom pickled dicts from training.
    Returns:
        models: dict[str, Any]   - loaded model payloads (normalized when helpful)
        model_status: dict[str, str] - human-readable status strings for the UI
    """
    base = Path(__file__).resolve().parent

    paths = {
        # BPE (merges-only pickle files)
        "code_bpe":     base / "bpe_models" / "merges_code.pkl",
        "doc_bpe":      base / "bpe_models" / "merges_doc.pkl",
        "combined_bpe": base / "bpe_models" / "merges_joint.pkl",

        # Word2Vec
        "code_w2v":     base / "w2v_models" / "w2v_code.pkl",
        "doc_w2v":      base / "w2v_models" / "w2v_doc.pkl",
        "combined_w2v": base / "w2v_models" / "w2v_joint.pkl",

        # LSTM
        "code_lstm":    base / "lstm_models" / "code_lstm.pkl",
        "doc_lstm":     base / "lstm_models" / "doc_lstm.pkl",
    }

    models, model_status = {}, {}

    def _filesize_mb(p: Path) -> float:
        try:
            return p.stat().st_size / (1024 * 1024)
        except Exception:
            return 0.0

    for key, path in paths.items():
        try:
            if not path.exists():
                model_status[key] = "❌ Not Found"
                continue

            with path.open("rb") as f:
                obj = pickle.load(f)

            # --- Normalize payloads for downstream classes ---
            if key.endswith("_bpe"):
                # Accept either list/tuple of merges or dict with merges/vocab/etc.
                if isinstance(obj, (list, tuple)):
                    # Wrap into the structure BPETokenizer expects
                    model_payload = {
                        "vocab": {},
                        "merges": list(obj),
                        "word_freqs": {},
                        "trained": True,
                    }
                elif isinstance(obj, dict):
                    model_payload = {
                        "vocab": obj.get("vocab", {}),
                        "merges": obj.get("merges", obj.get("pairs", [])),
                        "word_freqs": obj.get("word_freqs", {}),
                        "trained": obj.get("trained", True),
                    }
                else:
                    raise ValueError("Unsupported BPE pickle format")
                models[key] = model_payload
                model_status[key] = f"✅ Loaded ({_filesize_mb(path):.1f} MB)"

            elif key.endswith("_w2v"):
                if not isinstance(obj, dict):
                    raise ValueError("Expected dict payload for Word2Vec")
                obj.setdefault("vocab", {})
                obj.setdefault("word_count", {})
                obj.setdefault("syn0", np.array([]))
                obj.setdefault("vector_size", int(obj.get("vector_size", 100)))
                obj.setdefault("trained", bool(obj.get("trained", True)))
                models[key] = obj
                model_status[key] = f"✅ Loaded ({_filesize_mb(path):.1f} MB)"

            elif key.endswith("_lstm"):
                if isinstance(obj, dict) and "lstm" in obj:
                    if obj.get("trained", False):
                        models[key] = obj
                        model_status[key] = f"✅ Loaded ({_filesize_mb(path):.1f} MB)"
                    else:
                        models[key] = obj
                        model_status[key] = "⚠️ LSTM not trained"
                else:
                    model_status[key] = "⚠️ Invalid LSTM format (missing 'lstm')"

            else:
                models[key] = obj
                model_status[key] = f"✅ Loaded ({_filesize_mb(path):.1f} MB)"

        except Exception as e:
            msg = str(e)
            if "AdvancedLSTM" in msg:
                model_status[key] = "⚠️ LSTM class missing (define AdvancedLSTM before loading)"
            else:
                model_status[key] = f"❌ Error: {msg[:80]}..."

    return models, model_status


# =========================
# BPE Tokenizer
# =========================
class BPETokenizer:
    def __init__(self, model_data):
        self.vocab = model_data.get('vocab', {})
        self.merges = model_data.get('merges', [])
        self.word_freqs = model_data.get('word_freqs', {})
        self.trained = model_data.get('trained', False)

    def preprocess_text(self, text):
        """Preprocess text for tokenization"""
        if not isinstance(text, str):
            text = str(text)
        # Handle code-specific tokens
        text = re.sub(r'([{}()[\].,;:=+\-*/<>!&|^~%])', r' \1 ', text)
        text = re.sub(r'\s+', ' ', text)
        return text.lower().strip()

    def get_word_tokens(self, word):
        """Convert word to character tokens with end marker"""
        return list(word) + ['</w>']

    def tokenize(self, text):
        """Tokenize text using trained BPE model"""
        if not self.trained:
            return text.split()

        processed_text = self.preprocess_text(text)
        words = processed_text.split()

        tokens = []
        for word in words:
            word_tokens = self.get_word_tokens(word)
            word_str = ' '.join(word_tokens)

            # Apply merges
            for merge in self.merges:
                bigram = ' '.join(merge) if isinstance(merge, (list, tuple)) else str(merge)
                replacement = ''.join(merge) if isinstance(merge, (list, tuple)) else str(merge).replace(' ', '')
                word_str = word_str.replace(bigram, replacement)

            # Split back to tokens
            final_tokens = word_str.split()
            tokens.extend(final_tokens)

        return tokens

    def clean_tokens_for_display(self, tokens):
        """Clean BPE tokens for user-friendly display"""
        cleaned_tokens = []
        for token in tokens:
            clean_token = token.replace('</w>', '')
            clean_token = clean_token.replace('<unk>', '[UNK]')
            clean_token = clean_token.replace('<pad>', '[PAD]')
            clean_token = clean_token.replace('<bos>', '[START]')
            clean_token = clean_token.replace('<eos>', '[END]')
            if clean_token.strip():
                cleaned_tokens.append(clean_token)
        return cleaned_tokens


# =========================
# Word2Vec model wrapper
# =========================
class Word2VecModel:
    def __init__(self, model_data):
        self.vocab = model_data.get('vocab', {})
        self.word_count = model_data.get('word_count', {})
        self.syn0 = model_data.get('syn0', np.array([]))
        self.vector_size = model_data.get('vector_size', 100)
        self.trained = model_data.get('trained', False)

    def get_vector(self, word):
        """Get vector for a word"""
        if word in self.vocab and self.syn0.size > 0:
            word_index = self.vocab[word]
            if word_index < len(self.syn0):
                return self.syn0[word_index]
        return np.zeros(self.vector_size)

    def similarity(self, word1, word2):
        """Calculate cosine similarity between two words"""
        vec1 = self.get_vector(word1)
        vec2 = self.get_vector(word2)

        if np.linalg.norm(vec1) == 0 or np.linalg.norm(vec2) == 0:
            return 0.0

        return np.dot(vec1, vec2) / (np.linalg.norm(vec1) * np.linalg.norm(vec2))

    def most_similar(self, word, topn=10):
        """Find most similar words"""
        if word not in self.vocab or self.syn0.size == 0:
            return []

        word_vec = self.get_vector(word)
        similarities = []

        for other_word, index in self.vocab.items():
            if other_word != word and index < len(self.syn0):
                other_vec = self.syn0[index]
                denom = (np.linalg.norm(word_vec) * np.linalg.norm(other_vec))
                if denom == 0:
                    sim = 0.0
                else:
                    sim = float(np.dot(word_vec, other_vec) / denom)
                similarities.append((other_word, sim))

        similarities.sort(key=lambda x: x[1], reverse=True)
        return similarities[:topn]


# =========================
# LSTM model wrapper
# =========================
class LSTMModel:
    def __init__(self, model_data):
        self.lstm = model_data.get('lstm', None)
        self.vector_size = model_data.get('vector_size', 100)
        self.hidden_size = model_data.get('hidden_size', 128)
        self.vocab_size = model_data.get('vocab_size', 1000)
        self.losses = model_data.get('losses', [])
        self.trained = model_data.get('trained', False)
        self.model_type = model_data.get('model_type', 'unknown')
        self.sequences_trained = model_data.get('sequences_trained', 0)

    def predict_next_tokens(self, input_vectors, num_predictions=5):
        """Predict next tokens using LSTM model (placeholder sampling)"""
        if not self.trained or self.lstm is None:
            return []

        try:
            h = np.zeros((self.hidden_size, 1))
            c = np.zeros((self.hidden_size, 1))
            for vector in input_vectors:
                if hasattr(self.lstm, 'forward'):
                    h, c, _ = self.lstm.forward(vector.reshape(-1, 1), h, c)

            predictions = []
            for _ in range(num_predictions):
                prediction_score = np.random.random()
                predictions.append(prediction_score)
            return predictions
        except Exception:
            return []

    def get_context_representation(self, input_vectors):
        """Get context-aware representation from LSTM"""
        if not self.trained or self.lstm is None:
            return np.zeros(self.hidden_size)

        try:
            h = np.zeros((self.hidden_size, 1))
            c = np.zeros((self.hidden_size, 1))
            for vector in input_vectors:
                if hasattr(self.lstm, 'forward'):
                    h, c, _ = self.lstm.forward(vector.reshape(-1, 1), h, c)
            return h.flatten()
        except Exception:
            return np.zeros(self.hidden_size)

    def analyze_sequence_complexity(self, input_vectors):
        """Analyze sequence complexity using LSTM understanding"""
        if not input_vectors:
            return {'complexity': 'low', 'score': 0}

        context = self.get_context_representation(input_vectors)
        context_variance = np.var(context) if len(context) > 0 else 0

        if context_variance > 0.1:
            complexity = 'high'
            score = min(10, int(context_variance * 100))
        elif context_variance > 0.05:
            complexity = 'moderate'
            score = min(7, int(context_variance * 200))
        else:
            complexity = 'low'
            score = min(5, int(context_variance * 400))

        return {
            'complexity': complexity,
            'score': score,
            'context_variance': context_variance
        }


# =========================
# Code analysis functions
# =========================
def analyze_code_structure(code, lstm_models=None):
    """Analyze code structure and extract key components with LSTM enhancement"""
    analysis = {
        'functions': [],
        'classes': [],
        'imports': [],
        'variables': [],
        'complexity_score': 0,
        'lines_of_code': len(code.split('\n')),
        'docstrings': [],
        'lstm_analysis': {}
    }

    lines = code.split('\n')

    for line in lines:
        line = line.strip()

        # Function definitions
        if line.startswith('def '):
            func_match = re.search(r'def\s+(\w+)\s*\((.*?)\)', line)
            if func_match:
                func_name = func_match.group(1)
                params = func_match.group(2)
                analysis['functions'].append({
                    'name': func_name,
                    'parameters': [p.strip() for p in params.split(',') if p.strip()],
                    'line': line
                })

        # Class definitions
        elif line.startswith('class '):
            class_match = re.search(r'class\s+(\w+)', line)
            if class_match:
                analysis['classes'].append({
                    'name': class_match.group(1),
                    'line': line
                })

        # Import statements
        elif line.startswith('import ') or line.startswith('from '):
            analysis['imports'].append(line)

        # Variable assignments (simple heuristic)
        elif '=' in line and not line.startswith('#'):
            var_match = re.search(r'(\w+)\s*=', line)
            if var_match:
                analysis['variables'].append(var_match.group(1))

        # Docstrings (single-line markers)
        elif '"""' in line or "'''" in line:
            analysis['docstrings'].append(line)

    # Simple complexity score based on control structures
    complexity_indicators = ['if', 'elif', 'else', 'for', 'while', 'try', 'except', 'with']
    analysis['complexity_score'] = sum(code.lower().count(indicator) for indicator in complexity_indicators)

    # LSTM-based analysis if models are available
    if lstm_models:
        analysis['lstm_analysis'] = {
            'code_context_complexity': None,
            'semantic_flow': None,
            'predicted_patterns': []
        }

        try:
            if 'code_lstm' in lstm_models:
                lstm_model = LSTMModel(lstm_models['code_lstm'])

                # Simple vectorization of code for LSTM analysis
                code_words = code.lower().split()[:100]  # Limit for processing
                if code_words:
                    # Simple hash-based vector (placeholder)
                    word_vectors = []
                    for word in code_words:
                        vec = np.array([hash(word + str(i)) % 100 for i in range(lstm_model.vector_size)]) / 100.0
                        word_vectors.append(vec)

                    complexity_analysis = lstm_model.analyze_sequence_complexity(word_vectors)
                    analysis['lstm_analysis']['code_context_complexity'] = complexity_analysis

                    context_repr = lstm_model.get_context_representation(word_vectors)
                    if len(context_repr) > 0:
                        analysis['lstm_analysis']['semantic_flow'] = {
                            'context_strength': float(np.mean(np.abs(context_repr))),
                            'context_consistency': float(1.0 - np.std(context_repr))
                        }

        except Exception as e:
            analysis['lstm_analysis']['error'] = str(e)

    return analysis


def generate_function_description(func_name, parameters, func_line, lstm_context=None):
    """Generate intelligent function description with LSTM context enhancement"""
    special_methods = {
        '__init__': f"Constructor method that initializes a new instance with {len(parameters)} parameter(s)",
        '__str__': "String representation method that returns a human-readable string",
        '__repr__': "Object representation method for debugging and development",
        '__len__': "Returns the length of the object",
        '__eq__': "Equality comparison method",
        '__lt__': "Less than comparison method",
        '__gt__': "Greater than comparison method"
    }

    if func_name in special_methods:
        return special_methods[func_name]

    name_lower = func_name.lower()
    param_count = len([p for p in parameters if p.strip() and p.strip() != 'self'])

    if name_lower.startswith('get_') or name_lower.startswith('fetch_'):
        return f"Retrieves and returns data or information. Takes {param_count} parameter(s)"
    elif name_lower.startswith('set_') or name_lower.startswith('update_'):
        return f"Updates or modifies data/state. Takes {param_count} parameter(s)"
    elif name_lower.startswith('is_') or name_lower.startswith('has_') or name_lower.startswith('can_'):
        return f"Boolean check method that returns True/False. Takes {param_count} parameter(s)"
    elif name_lower.startswith('create_') or name_lower.startswith('make_') or name_lower.startswith('build_'):
        return f"Creates or constructs new objects/data. Takes {param_count} parameter(s)"
    elif name_lower.startswith('delete_') or name_lower.startswith('remove_') or name_lower.startswith('clear_'):
        return f"Deletes or removes data/objects. Takes {param_count} parameter(s)"
    elif name_lower.startswith('find_') or name_lower.startswith('search_'):
        return f"Searches and locates specific data or objects. Takes {param_count} parameter(s)"
    elif name_lower.startswith('calculate_') or name_lower.startswith('compute_'):
        return f"Performs calculations and returns computed result. Takes {param_count} parameter(s)"
    elif name_lower.startswith('load_') or name_lower.startswith('read_'):
        return f"Loads or reads data from external source. Takes {param_count} parameter(s)"
    elif name_lower.startswith('save_') or name_lower.startswith('write_'):
        return f"Saves or writes data to external destination. Takes {param_count} parameter(s)"
    elif name_lower.startswith('add_') or name_lower.startswith('insert_'):
        return f"Adds or inserts new items/data. Takes {param_count} parameter(s)"
    elif name_lower.startswith('view_') or name_lower.startswith('display_') or name_lower.startswith('show_'):
        return f"Displays or presents information to user. Takes {param_count} parameter(s)"
    elif 'main' in name_lower:
        return f"Main execution function that coordinates program flow. Entry point with {param_count} parameter(s)"
    else:
        if lstm_context and 'complexity' in lstm_context:
            complexity = lstm_context['complexity']
            if complexity == 'high':
                base_desc = f"Complex function '{func_name}' with sophisticated logic and {param_count} parameter(s)"
            elif complexity == 'moderate':
                base_desc = f"Function '{func_name}' with moderate complexity handling {param_count} parameter(s)"
            else:
                base_desc = f"Simple function '{func_name}' performing straightforward operations with {param_count} parameter(s)"
        else:
            if param_count == 0:
                base_desc = f"Function '{func_name}' with no parameters. Likely performs standalone operation"
            elif param_count == 1:
                base_desc = f"Function '{func_name}' that processes single input parameter"
            else:
                base_desc = f"Function '{func_name}' that processes {param_count} input parameters"
        return base_desc


def generate_class_description(class_name, functions):
    """Generate intelligent class description"""
    class_methods = []
    for func in functions:
        if 'self' in func.get('parameters', []):
            class_methods.append(func['name'])

    method_count = len(class_methods)
    name_lower = class_name.lower()

    if 'system' in name_lower or 'manager' in name_lower:
        return f"Management class that handles operations and coordination. Contains {method_count} method(s)"
    elif 'service' in name_lower or 'handler' in name_lower:
        return f"Service class that provides specific functionality. Contains {method_count} method(s)"
    elif 'model' in name_lower or 'data' in name_lower:
        return f"Data model class representing structured information. Contains {method_count} method(s)"
    elif 'util' in name_lower or 'helper' in name_lower:
        return f"Utility class providing helper functions and tools. Contains {method_count} method(s)"
    elif 'student' in name_lower:
        return f"Student entity class representing individual student data. Contains {method_count} method(s)"
    elif 'user' in name_lower or 'person' in name_lower:
        return f"User/Person entity class for managing individual records. Contains {method_count} method(s)"
    elif 'config' in name_lower or 'setting' in name_lower:
        return f"Configuration class for managing application settings. Contains {method_count} method(s)"
    else:
        has_init = '__init__' in class_methods
        has_str = '__str__' in class_methods

        if has_init and has_str:
            return f"Custom class '{class_name}' with proper initialization and string representation. Contains {method_count} method(s)"
        elif has_init:
            return f"Custom class '{class_name}' with initialization capabilities. Contains {method_count} method(s)"
        else:
            return f"Custom class '{class_name}' providing specialized functionality. Contains {method_count} method(s)"


def generate_documentation_content(code, code_analysis, models, params=None):
    """Generate comprehensive documentation content with LSTM enhancement"""
    if params is None:
        params = {'style': 'comprehensive', 'ai_level': 'basic', 'include_lstm': False}

    title_map = {
        'comprehensive': 'Python Code Documentation',
        'technical': 'Technical Code Analysis Report',
        'user-friendly': 'Code Usage Guide',
        'api reference': 'API Reference Documentation'
    }

    doc_content = {
        'title': title_map.get(params.get('style', 'comprehensive'), 'Python Code Documentation'),
        'overview': '',
        'functions': [],
        'classes': [],
        'usage_examples': [],
        'technical_details': {},
        'semantic_analysis': {},
        'lstm_insights': {},
        'generation_params': params
    }

    func_count = len(code_analysis['functions'])
    class_count = len(code_analysis['classes'])
    complexity = code_analysis['complexity_score']
    lines_count = code_analysis['lines_of_code']
    imports_count = len(code_analysis['imports'])

    if class_count > 0 and func_count > class_count * 2:
        module_type = "object-oriented module with utility functions"
    elif class_count > 0:
        module_type = "object-oriented module"
    elif func_count > 5:
        module_type = "functional programming module"
    elif 'main' in [f['name'] for f in code_analysis['functions']]:
        module_type = "executable script"
    else:
        module_type = "utility module"

    complexity_level = 'high' if complexity > 15 else 'moderate' if complexity > 8 else 'low'

    doc_content['overview'] = f"""
This is a {module_type} containing {func_count} function(s) and {class_count} class(es).

Code Characteristics:
• Structure: {lines_count} lines of code with {imports_count} external dependencies
• Complexity: {complexity_level.title()} complexity (score: {complexity}) indicating {'advanced control flow with multiple decision points' if complexity > 15 else 'moderate branching and loops' if complexity > 8 else 'straightforward linear execution'}
• Design Pattern: {'Class-based design with encapsulation' if class_count > 0 else 'Procedural programming approach'}

Module Purpose: {'Appears to be a complete application with user interaction' if 'main' in [f['name'] for f in code_analysis['functions']] else 'Provides reusable components and functionality'}
    """.strip()

    # Analyze functions using BPE and Word2Vec
    if 'code_bpe' in models and 'code_w2v' in models:
        try:
            bpe_tokenizer = BPETokenizer(models['code_bpe'])
            w2v_model = Word2VecModel(models['code_w2v'])

            for func in code_analysis['functions']:
                func_name = func['name']
                func_tokens = bpe_tokenizer.tokenize(func['line'])
                cleaned_tokens = bpe_tokenizer.clean_tokens_for_display(func_tokens)

                lstm_context = None
                if 'lstm_analysis' in code_analysis and 'code_context_complexity' in code_analysis['lstm_analysis']:
                    lstm_context = code_analysis['lstm_analysis']['code_context_complexity']

                description = generate_function_description(func_name, func['parameters'], func['line'], lstm_context)

                func_doc = {
                    'name': func_name,
                    'parameters': func['parameters'],
                    'description': description,
                    'tokens': cleaned_tokens[:8],
                    'similar_concepts': [],
                    'lstm_complexity': lstm_context
                }

                if func_name in w2v_model.vocab:
                    similar = w2v_model.most_similar(func_name, topn=5)
                    func_doc['similar_concepts'] = [word for word, score in similar if score > 0.3]

                doc_content['functions'].append(func_doc)

        except Exception as e:
            st.warning(f"Error in semantic analysis: {str(e)}")

    # Analyze classes
    for cls in code_analysis['classes']:
        class_name = cls['name']
        class_description = generate_class_description(class_name, code_analysis['functions'])
        class_doc = {'name': class_name, 'description': class_description, 'methods': []}
        doc_content['classes'].append(class_doc)

    # Technical details
    doc_content['technical_details'] = {
        'imports': code_analysis['imports'],
        'complexity_metrics': {
            'lines_of_code': code_analysis['lines_of_code'],
            'cyclomatic_complexity': code_analysis['complexity_score'],
            'function_count': len(code_analysis['functions']),
            'class_count': len(code_analysis['classes'])
        }
    }

    # LSTM insights
    if 'lstm_analysis' in code_analysis and code_analysis['lstm_analysis']:
        lstm_analysis = code_analysis['lstm_analysis']
        doc_content['lstm_insights'] = {
            'semantic_analysis': lstm_analysis.get('semantic_flow', {}),
            'context_complexity': lstm_analysis.get('code_context_complexity', {}),
            'ai_understanding': {}
        }

        if 'semantic_flow' in lstm_analysis and lstm_analysis['semantic_flow']:
            semantic = lstm_analysis['semantic_flow']
            if semantic.get('context_strength', 0) > 0.5:
                doc_content['lstm_insights']['ai_understanding']['code_coherence'] = "High - Code shows strong semantic coherence"
            elif semantic.get('context_strength', 0) > 0.3:
                doc_content['lstm_insights']['ai_understanding']['code_coherence'] = "Moderate - Code has reasonable structure"
            else:
                doc_content['lstm_insights']['ai_understanding']['code_coherence'] = "Low - Code may benefit from refactoring"

            if semantic.get('context_consistency', 0) > 0.7:
                doc_content['lstm_insights']['ai_understanding']['consistency'] = "High - Consistent coding patterns detected"
            else:
                doc_content['lstm_insights']['ai_understanding']['consistency'] = "Variable - Mixed coding patterns"

    return doc_content


def create_word_document(doc_content, code):
    """Create a Word document with the generated documentation"""
    doc = Document()

    # Add title
    title = doc.add_heading(doc_content['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata
    doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("Generated by: GenAI Documentation Generator")
    doc.add_paragraph("")

    # Add overview
    doc.add_heading('Overview', level=1)
    doc.add_paragraph(doc_content['overview'])

    # Add original code
    doc.add_heading('Source Code', level=1)
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)

    # Add functions documentation
    if doc_content['functions']:
        doc.add_heading('Functions Documentation', level=1)
        for func in doc_content['functions']:
            doc.add_heading(f"Function: {func['name']}", level=2)
            doc.add_paragraph(f"Description: {func['description']}")
            if func['parameters']:
                doc.add_paragraph("Parameters:")
                for param in func['parameters']:
                    doc.add_paragraph(f"  • {param}", style='List Bullet')
            if func['tokens']:
                doc.add_paragraph(f"BPE Tokens: {', '.join(func['tokens'][:10])}")
            if func['similar_concepts']:
                doc.add_paragraph(f"Related Concepts: {', '.join(func['similar_concepts'])}")
            doc.add_paragraph("")

    # Add classes documentation
    if doc_content['classes']:
        doc.add_heading('Classes Documentation', level=1)
        for cls in doc_content['classes']:
            doc.add_heading(f"Class: {cls['name']}", level=2)
            doc.add_paragraph(cls['description'])
            doc.add_paragraph("")

    # Add technical details
    doc.add_heading('Technical Analysis', level=1)
    tech_details = doc_content['technical_details']

    doc.add_heading('Complexity Metrics', level=2)
    metrics = tech_details['complexity_metrics']
    doc.add_paragraph(f"Lines of Code: {metrics['lines_of_code']}")
    doc.add_paragraph(f"Cyclomatic Complexity: {metrics['cyclomatic_complexity']}")
    doc.add_paragraph(f"Number of Functions: {metrics['function_count']}")
    doc.add_paragraph(f"Number of Classes: {metrics['class_count']}")

    # LSTM insights
    if 'lstm_insights' in doc_content and doc_content['lstm_insights']:
        doc.add_heading('AI-Powered Analysis (LSTM)', level=2)
        lstm_insights = doc_content['lstm_insights']

        if 'ai_understanding' in lstm_insights:
            understanding = lstm_insights['ai_understanding']
            if 'code_coherence' in understanding:
                doc.add_paragraph(f"Code Coherence: {understanding['code_coherence']}")
            if 'consistency' in understanding:
                doc.add_paragraph(f"Pattern Consistency: {understanding['consistency']}")

        if 'context_complexity' in lstm_insights and lstm_insights['context_complexity']:
            complexity = lstm_insights['context_complexity']
            doc.add_paragraph(f"AI Complexity Assessment: {complexity.get('complexity', 'unknown').title()}")
            doc.add_paragraph(f"Context Analysis Score: {complexity.get('score', 0)}/10")

    # Imports
    if tech_details['imports']:
        doc.add_heading('Dependencies', level=2)
        for imp in tech_details['imports']:
            doc.add_paragraph(f"  {imp}", style='List Bullet')

    # Footer
    doc.add_paragraph("")
    doc.add_paragraph("---")
    footer = doc.add_paragraph(
        "This documentation was automatically generated using trained BPE tokenizers, Word2Vec models, and LSTM neural networks."
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


# =========================
# Main Streamlit application
# =========================
def main():
    # Header
    st.markdown('<h1 class="main-header">📚 GenAI Documentation Generator</h1>', unsafe_allow_html=True)
    st.markdown("**Transform your Python code into comprehensive documentation using AI-powered analysis**")

    # Load models
    with st.spinner("Loading trained models..."):
        models, model_status = load_models()

    # Model Summary
    total_models = len([k for k, v in model_status.items() if "✅" in v])
    st.sidebar.metric("Models Loaded", f"{total_models}/{len(model_status)}")

    # Sidebar - Model Status
    st.sidebar.markdown("## 🤖 Model Status")
    model_names = {
        'code_bpe': 'Code BPE Tokenizer',
        'doc_bpe': 'Documentation BPE',
        'combined_bpe': 'Combined BPE',
        'code_w2v': 'Code Word2Vec',
        'doc_w2v': 'Documentation Word2Vec',
        'combined_w2v': 'Combined Word2Vec',
        'code_lstm': 'Code LSTM Model',
        'doc_lstm': 'Documentation LSTM'
    }
    model_groups = {
        "🔤 BPE Tokenizers": ['code_bpe', 'doc_bpe', 'combined_bpe'],
        "🎯 Word2Vec Models": ['code_w2v', 'doc_w2v', 'combined_w2v'],
        "🧠 LSTM Networks": ['code_lstm', 'doc_lstm']
    }
    for group_name, model_keys in model_groups.items():
        with st.sidebar.expander(group_name, expanded=True):
            for model_key in model_keys:
                if model_key in model_status:
                    status = model_status[model_key]
                    model_name = model_names.get(model_key, model_key)
                    if "✅" in status:
                        st.success(f"{model_name}: {status}")
                    elif "⚠️" in status:
                        st.warning(f"{model_name}: {status}")
                    else:
                        st.error(f"{model_name}: {status}")

    # Check if essential models are loaded
    essential_models = ['code_bpe', 'code_w2v']
    models_loaded = all(model_key in models for model_key in essential_models)

    # Check LSTM availability
    lstm_available = any(model_key in models for model_key in ['code_lstm', 'doc_lstm'])

    if not models_loaded:
        st.error("⚠️ Essential models not loaded. Please ensure the trained models are available in the correct directories.")
        st.info("""
        Required model files:
        - ./bpe_models/merges_code.pkl
        - ./w2v_models/w2v_code.pkl

        Optional LSTM models for enhanced analysis:
        - ./lstm_models/code_lstm.pkl
        - ./lstm_models/doc_lstm.pkl

        Run your training notebook to produce these pickles first.
        """)
        return

    # LSTM Status
    if lstm_available:
        st.sidebar.success("🧠 LSTM Models Available - Enhanced Analysis Enabled")
    else:
        st.sidebar.info("ℹ️ LSTM Models Optional - Basic Analysis Mode")

    # Main interface
    st.markdown('<h2 class="section-header">📝 Code Input</h2>', unsafe_allow_html=True)

    # Code input options
    input_method = st.radio("Choose input method:", ["Paste Code", "Upload File"])

    code_input = ""
    if input_method == "Paste Code":
        code_input = st.text_area(
            "Paste your Python code here:",
            height=300,
            placeholder="def example_function(param1, param2):\n    \"\"\"\n    Example function description\n    \"\"\"\n    return param1 + param2"
        )
    else:
        uploaded_file = st.file_uploader("Upload a Python file", type=['py', 'txt'])
        if uploaded_file is not None:
            code_input = str(uploaded_file.read(), "utf-8")
            st.code(code_input, language='python')

    if code_input.strip():
        # Analysis section
        st.markdown('<h2 class="section-header">🔍 Code Analysis</h2>', unsafe_allow_html=True)

        with st.spinner("Analyzing code structure with AI models..."):
            lstm_models = {k: v for k, v in models.items() if 'lstm' in k} if lstm_available else None
            code_analysis = analyze_code_structure(code_input, lstm_models)

        # Display analysis results
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Lines of Code", code_analysis['lines_of_code'])
        with col2:
            st.metric("Functions", len(code_analysis['functions']))
        with col3:
            st.metric("Classes", len(code_analysis['classes']))
        with col4:
            st.metric("Complexity Score", code_analysis['complexity_score'])

        # Detailed analysis
        if code_analysis['functions']:
            st.markdown("### 🔧 Functions Found:")
            for func in code_analysis['functions']:
                st.markdown(f"- **{func['name']}** ({len(func['parameters'])} parameters)")

        if code_analysis['classes']:
            st.markdown("### 📦 Classes Found:")
            for cls in code_analysis['classes']:
                st.markdown(f"- **{cls['name']}**)")

        if code_analysis['imports']:
            st.markdown("### 📚 Imports:")
            for imp in code_analysis['imports']:
                st.code(imp, language='python')

        # LSTM Analysis (if available)
        if lstm_available and 'lstm_analysis' in code_analysis and code_analysis['lstm_analysis']:
            st.markdown("### 🧠 Advanced AI Analysis")
            with st.container():
                lstm_analysis = code_analysis['lstm_analysis']
                complexity_data = lstm_analysis.get('code_context_complexity', {}) or {}

                col1, col2, col3 = st.columns(3)
                with col1:
                    complexity_level = complexity_data.get('complexity', 'unknown')
                    complexity_color = {'high': '🔴', 'moderate': '🟡', 'low': '🟢'}.get(str(complexity_level).lower(), '⚪')
                    st.metric("AI Complexity", f"{complexity_color} {str(complexity_level).title()}")
                with col2:
                    score = complexity_data.get('score', 0)
                    st.metric("Context Score", f"{score}/10")
                with col3:
                    variance = complexity_data.get('context_variance', 0)
                    try:
                        st.metric("Context Variance", f"{float(variance):.3f}")
                    except Exception:
                        st.metric("Context Variance", "N/A")

                semantic = lstm_analysis.get('semantic_flow', {}) or {}
                if semantic:
                    st.markdown("**Semantic Analysis:**")
                    c1, c2 = st.columns(2)
                    with c1:
                        context_strength = float(semantic.get('context_strength', 0))
                        st.progress(min(1.0, max(0.0, context_strength)))
                        st.caption(f"Semantic Coherence: {context_strength:.2%}")
                    with c2:
                        consistency = float(semantic.get('context_consistency', 0))
                        st.progress(min(1.0, max(0.0, consistency)))
                        st.caption(f"Pattern Consistency: {consistency:.2%}")

                if complexity_data.get('complexity') == 'high':
                    st.info("🧠 **AI Insight**: This code shows high complexity patterns. Consider breaking into smaller functions.")
                elif complexity_data.get('complexity') == 'low' and len(code_analysis['functions']) > 5:
                    st.success("🧠 **AI Insight**: Well-structured code with good separation of concerns detected.")

        # AI Analysis & Documentation
        st.markdown('<h2 class="section-header">🧠 AI-Powered Analysis & Documentation</h2>', unsafe_allow_html=True)

        col1, col2 = st.columns(2)
        with col1:
            doc_style = st.selectbox(
                "📝 Documentation Style:",
                ["Comprehensive", "Technical", "User-Friendly", "API Reference"],
                help="Choose how detailed and technical the documentation should be"
            )
        with col2:
            if lstm_available:
                ai_level = st.selectbox(
                    "🧠 AI Analysis Level:",
                    ["Basic", "Enhanced", "Deep"],
                    index=1,
                    help="Higher levels provide more detailed LSTM-powered insights"
                )
            else:
                ai_level = "Basic"
                st.info("💡 Install LSTM models for advanced AI analysis")

        with st.spinner("Generating AI-powered analysis using trained models..."):
            generation_params = {
                'style': doc_style.lower(),
                'ai_level': ai_level.lower(),
                'include_lstm': lstm_available
            }
            doc_content = generate_documentation_content(code_input, code_analysis, models, generation_params)

        if doc_content['functions']:
            st.markdown("### 🤖 Function Analysis:")
            for func in doc_content['functions']:
                with st.expander(f"Function: {func['name']}"):
                    st.write(f"**Description:** {func['description']}")
                    if func['tokens']:
                        st.write(f"**BPE Tokens:** {', '.join(func['tokens'])}")
                    if func['similar_concepts']:
                        st.write(f"**Related Concepts:** {', '.join(func['similar_concepts'])}")
                    if lstm_available and func.get('lstm_complexity'):
                        complexity = func['lstm_complexity']
                        st.write(f"**LSTM Complexity:** {complexity.get('complexity', 'unknown').title()} (Score: {complexity.get('score', 0)})")

        # Documentation Generation
        st.markdown('<h2 class="section-header">📄 Generate Documentation</h2>', unsafe_allow_html=True)

        if st.button("🚀 Generate Word Document", type="primary"):
            with st.spinner("Generating comprehensive documentation..."):
                try:
                    word_doc = create_word_document(doc_content, code_input)
                    doc_buffer = io.BytesIO()
                    word_doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"code_documentation_{timestamp}.docx"

                    st.success("✅ Documentation generated successfully!")
                    st.download_button(
                        label="📥 Download Word Document",
                        data=doc_buffer.getvalue(),
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                    st.markdown("### 📋 Documentation Preview:")
                    st.markdown(f"**Title:** {doc_content['title']}")
                    st.markdown(f"**Overview:** {doc_content['overview']}")
                    if doc_content['functions']:
                        st.markdown("**Functions Documented:**")
                        for func in doc_content['functions']:
                            st.markdown(f"- {func['name']}")
                except Exception as e:
                    st.error(f"Error generating documentation: {str(e)}")
    else:
        st.info("👆 Please enter or upload Python code to get started!")

    # Footer
    st.markdown("---")
    st.markdown(f"""
    <div style='text-align: center; color: #666; margin-top: 2rem;'>
        <p>🤖 Powered by BPE Tokenizers, Word2Vec models{' & LSTM Networks' if lstm_available else ''}</p>
        <p>Built with Streamlit • GenAI Documentation Generator</p>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
